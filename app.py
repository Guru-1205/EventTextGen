from flask import Flask, render_template, request, send_file
from PIL import Image, ImageDraw, ImageFont
from io import BytesIO
from docx import Document
from docx.shared import Inches
import os
import google.generativeai as genai
from dotenv import load_dotenv

load_dotenv()
app = Flask(__name__)

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

doc_stream = None
genai.configure(api_key=os.getenv('GENAI_API_KEY'))


def generate_text(prompt):
    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(
        max_output_tokens=(100*3),
    ))
    return response.text


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/')
def index():
    return render_template('index1.html', document_url=None)


@app.route('/upload', methods=['POST'])
def upload():
    def textsize(text, font):
        im = Image.new(mode="P", size=(0, 0))
        draw = ImageDraw.Draw(im)
        _, _, width, height = draw.textbbox((0, 0), text=text, font=font)
        return width, height
    global doc_stream  # Allow modification of global variable

    # Check if app.root_path is None
    if app.root_path is None:
        # Provide a fallback path or handle the error appropriately
        root_path = os.getcwd()  # Example fallback to current working directory
    else:
        root_path = app.root_path

    files = request.files.getlist('files')
    edited_images = []
    error_msgs = []

    for file in files:
        if file and allowed_file(file.filename):
            try:
                img = Image.open(file)
                img = img.resize((1027, 515))
                draw = ImageDraw.Draw(img)
                font_size = 20

                # Correct path to the font file
                font_path = os.path.join(root_path, 'fonts', 'ARIAL.TTF')
                font = ImageFont.truetype(font_path, font_size)

                event_name = request.form.get('event_name', 'Your Event Name')

                # Calculate the size of the text
                text_width, text_height = textsize(event_name, font)

                # Calculate the rectangle coordinates to center it at the top with a small gap
                rect_height = 34
                gap = 30  # Adjust the gap as needed
                rect_width = text_width + 2 * gap  # Add some padding
                rect_coordinates = [img.width // 2 - rect_width // 2, gap,
                                    img.width // 2 + rect_width // 2, gap + rect_height]

                # Draw the white rectangle
                draw.rectangle(rect_coordinates, fill=(255, 255, 255))

                # Calculate the text position to center it within the rectangle
                text_position = ((img.width - text_width) // 2,
                                 gap + (rect_height - text_height) // 2)

                # Draw the event name text
                draw.text(text_position, event_name, font=font, fill=(0, 0, 0))

                edited_image_stream = BytesIO()
                img.save(edited_image_stream, format='PNG')
                edited_image_stream.seek(0)

                edited_images.append(edited_image_stream)

            except IOError as e:
                error_msgs.append(
                    f'Error processing image {file.filename}: {str(e)}')

        else:
            error_msgs.append(f'Invalid file format: {file.filename}')

    if error_msgs:
        return render_template('index1.html', error='\n'.join(error_msgs))

    context = request.form.get('context', 'Your event context here.')

    document = Document()
    event_name = request.form.get('event_name', 'Your Event Name')
    document.add_heading(event_name, level=1)

    for edited_image_stream in edited_images:
        document.add_picture(edited_image_stream, width=Inches(4.0))

    rephrased_context = generate_text(context)
    document.add_paragraph('Rephrased Context:')
    document.add_paragraph(rephrased_context)
    doc_stream = BytesIO()
    document.save(doc_stream)
    doc_stream.seek(0)

    return send_file(doc_stream,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, attachment_filename='output_document.docx')


if __name__ == '__main__':
    app.run(debug=True)
